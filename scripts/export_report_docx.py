"""
Export BAO_CAO_DU_AN.md → BAO_CAO_DU_AN.docx (Word)

Dùng python-docx để tạo Word document với:
- Heading styles
- Tables
- Code blocks
- Embedded images
- Page breaks
"""

from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR    = Path(__file__).parent.parent
MD_FILE     = BASE_DIR / "BAO_CAO_DU_AN.md"
OUT_FILE    = BASE_DIR / "BAO_CAO_DU_AN.docx"
IMG_DIR     = BASE_DIR / "docs" / "screenshots"

# ── Screenshot captions ──────────────────────────────────────────────────────
SCREENSHOTS = [
    ("01_main_page.png",          "Hình 1: Giao diện chính TaxAI — Trang chủ sau khi khởi động"),
    ("02_sidebar_area.png",       "Hình 2: Sidebar — Bộ lọc văn bản, toggle tùy chọn, lịch sử hội thoại"),
    ("03_chat_area.png",          "Hình 3: Khu vực chat chính — Tiêu đề và câu hỏi gợi ý"),
    ("04_suggested_questions.png","Hình 4: 6 câu hỏi gợi ý theo chủ đề thực tế"),
    ("05_question_in_input.png",  "Hình 5: Câu hỏi tự động điền sau khi click gợi ý — đang xử lý"),
    ("06_full_ui.png",            "Hình 6: Toàn bộ giao diện — bố cục sidebar + chat area"),
]


def set_cell_bg(cell, hex_color: str):
    """Set background color for a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 0:
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)   # blue
    return h


def add_code_block(doc: Document, code: str):
    """Add a code block with monospace font and grey background."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)

    # Add grey shading to paragraph
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F3F4F6")
    pPr.append(shd)

    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_table_from_md(doc: Document, lines: list[str]):
    """Parse markdown table and add to doc."""
    rows = []
    for line in lines:
        if line.strip().startswith("|---") or re.match(r"^\|[-|\s]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = row.cells[j]
                # Strip markdown bold
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                clean = re.sub(r"`(.*?)`", r"\1", clean)
                cell.text = clean
                if i == 0:
                    # Header row: grey background + bold
                    set_cell_bg(cell, "E5E7EB")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                    cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else None
                    p = cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].bold = True


def add_image_section(doc: Document, filename: str, caption: str):
    """Add screenshot image with caption."""
    img_path = IMG_DIR / filename
    if not img_path.exists():
        doc.add_paragraph(f"[Hình ảnh không tìm thấy: {filename}]")
        return

    try:
        doc.add_picture(str(img_path), width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        for run in cap.runs:
            run.italic  = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    except Exception as e:
        doc.add_paragraph(f"[Lỗi thêm ảnh {filename}: {e}]")


def process_inline(text: str):
    """Return list of (text, bold, italic, code) tuples from inline markdown."""
    parts = []
    # Simple tokeniser for **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False, False))
        tok = m.group()
        if tok.startswith("**"):
            parts.append((tok[2:-2], True, False, False))
        elif tok.startswith("*"):
            parts.append((tok[1:-1], False, True, False))
        elif tok.startswith("`"):
            parts.append((tok[1:-1], False, False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False, False))
    return parts


def add_rich_paragraph(doc: Document, text: str, style=None, indent_cm=0.0):
    """Add paragraph with inline bold/italic/code."""
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)

    for chunk, bold, italic, code in process_inline(text):
        run = para.add_run(chunk)
        run.bold   = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)


def build_docx():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Read markdown ─────────────────────────────────────────────────────────
    md = MD_FILE.read_text(encoding="utf-8")
    lines = md.splitlines()

    in_code  = False
    code_buf: list[str] = []
    in_table = False
    table_buf: list[str] = []

    # Track where to inject screenshots (after UI section header)
    ui_section_done = False
    screenshot_injected = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Code block ────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf  = []
                in_code   = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            if not in_table:
                in_table  = True
                table_buf = []
            table_buf.append(line)
            i += 1
            continue
        else:
            if in_table:
                add_table_from_md(doc, table_buf)
                doc.add_paragraph()
                in_table  = False
                table_buf = []

        stripped = line.strip()

        # ── Horizontal rule ───────────────────────────────────────────────────
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            add_heading(doc, text, level)

            # Inject screenshots right after the UI section heading
            if "GIAO DIỆN NGƯỜI DÙNG" in text.upper() and not screenshot_injected:
                ui_section_done = True

            i += 1
            continue

        # ── Screenshot injection (after "Mô tả từng tính năng UI" subsection) ─
        if ui_section_done and not screenshot_injected and "6.3 Screenshot" in line:
            # Add all screenshots here
            doc.add_heading("Ảnh chụp màn hình UI", level=3)
            for fname, caption in SCREENSHOTS:
                add_image_section(doc, fname, caption)
            screenshot_injected = True

        # ── Bullet points ─────────────────────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            add_rich_paragraph(doc, text, style="List Bullet", indent_cm=0.0)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        m2 = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m2:
            text = m2.group(2)
            add_rich_paragraph(doc, text, style="List Number")
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if not stripped:
            if not (i > 0 and not lines[i-1].strip()):  # avoid double spacing
                doc.add_paragraph()
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        add_rich_paragraph(doc, stripped)
        i += 1

    # Flush any remaining table/code
    if in_table:
        add_table_from_md(doc, table_buf)
    if in_code:
        add_code_block(doc, "\n".join(code_buf))

    # ── If screenshots not yet injected, add at end of doc ────────────────────
    if not screenshot_injected:
        doc.add_page_break()
        doc.add_heading("Ảnh chụp màn hình Giao diện", level=2)
        for fname, caption in SCREENSHOTS:
            add_image_section(doc, fname, caption)

    doc.save(str(OUT_FILE))
    print(f"✅ Exported: {OUT_FILE}")
    print(f"   Pages: ~{len(doc.paragraphs)//40 + 1} (ước tính)")


if __name__ == "__main__":
    build_docx()
