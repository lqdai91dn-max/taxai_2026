"""
Convert BAO_CAO_DU_AN_v2.md → BAO_CAO_DU_AN_v2.docx
Định dạng báo cáo chuyên nghiệp chuẩn in ấn (A4, Times New Roman, VN)
"""
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Constants ────────────────────────────────────────────────────
FONT_MAIN = "Times New Roman"
FONT_CODE = "Courier New"
SIZE_BODY  = Pt(13)
SIZE_H1    = Pt(16)
SIZE_H2    = Pt(14)
SIZE_H3    = Pt(13)
SIZE_CODE  = Pt(10)
SIZE_TABLE = Pt(11)
SIZE_SMALL = Pt(10)

COLOR_H1     = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
COLOR_H2     = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
COLOR_H3     = RGBColor(0x2F, 0x54, 0x96)   # blue
COLOR_TABLE_HEADER = RGBColor(0x2E, 0x74, 0xB5)
COLOR_CODE_BG      = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_ACCENT       = RGBColor(0x70, 0xAD, 0x47)  # green accent

SRC = Path("BAO_CAO_DU_AN_v2.md")
DST = Path("BAO_CAO_DU_AN_v2.docx")


# ─── Document setup ────────────────────────────────────────────────
def create_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_MAIN
    style.font.size = SIZE_BODY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.3

    return doc


def set_font_vn(run, bold=False, italic=False, size=None, color=None, name=None):
    run.font.name = name or FONT_MAIN
    run.font.bold = bold
    run.font.italic = italic
    if size: run.font.size = size
    if color: run.font.color.rgb = color
    # Force Vietnamese font fallback
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    fn = name or FONT_MAIN
    rFonts.set(qn('w:ascii'), fn)
    rFonts.set(qn('w:hAnsi'), fn)
    rFonts.set(qn('w:cs'), fn)
    existing = rpr.find(qn('w:rFonts'))
    if existing is not None:
        rpr.remove(existing)
    rpr.insert(0, rFonts)


def shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_cell_border(cell):
    """Add thin borders to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'BFBFBF')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_number(doc):
    """Add 'Trang X / Y' to footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = FONT_MAIN
    run.font.size = Pt(10)
    run.text = "Trang "
    # PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run2 = p.add_run(" / ")
    run2.font.name = FONT_MAIN
    run2.font.size = Pt(10)
    # NUMPAGES field
    run3 = p.add_run()
    run3.font.name = FONT_MAIN
    run3.font.size = Pt(10)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar3)
    run3._r.append(instrText2)
    run3._r.append(fldChar4)


# ─── Cover page ────────────────────────────────────────────────────
def add_cover(doc: Document):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO DỰ ÁN")
    run.font.name = FONT_MAIN
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_H1
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TaxAI")
    run.font.name = FONT_MAIN
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_H2
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hệ thống Tư vấn Thuế Tự động")
    run.font.name = FONT_MAIN
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = COLOR_H3
    p.paragraph_format.space_after = Pt(36)

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.name = FONT_MAIN
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_H2
    p.paragraph_format.space_after = Pt(24)

    # Metadata table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    meta = [
        ("Phiên bản",        "2.0"),
        ("Ngày cập nhật",    "10/04/2026"),
        ("Trạng thái",       "Production-ready — 100% pass rate (225/225)"),
        ("Công nghệ chính",  "Gemini Flash + ChromaDB + Vietnamese-SBERT"),
        ("Corpus",           "20 văn bản pháp luật | 225 câu benchmark"),
    ]
    col_widths = [Cm(5), Cm(11)]
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        shade_cell(row.cells[0], "2E74B5")
        shade_cell(row.cells[1], "DEEAF1")
        for cell in row.cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(k)
        set_font_vn(r0, bold=True, size=Pt(12), color=RGBColor(0xFF, 0xFF, 0xFF))
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(v)
        set_font_vn(r1, size=Pt(12), color=COLOR_H1)

    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()


# ─── Heading helpers ───────────────────────────────────────────────
def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font_vn(run, bold=True, size=SIZE_H1, color=COLOR_H1)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font_vn(run, bold=True, size=SIZE_H2, color=COLOR_H2)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font_vn(run, bold=True, size=SIZE_H3, color=COLOR_H3)
    return p


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font_vn(run, bold=True, size=SIZE_BODY, color=RGBColor(0x40, 0x40, 0x40))
    return p


# ─── Inline formatting ─────────────────────────────────────────────
def add_inline(paragraph, text):
    """Parse **bold**, `code`, plain text within a paragraph."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font_vn(run, bold=True, size=SIZE_BODY)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_CODE
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            run = paragraph.add_run(part)
            set_font_vn(run, size=SIZE_BODY)


def add_body(doc, text):
    """Add a normal paragraph with inline formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    add_inline(p, text)
    return p


# ─── Table parser ──────────────────────────────────────────────────
def is_separator_row(row_text):
    return bool(re.match(r'^\|[-| :]+\|$', row_text.strip()))


def parse_md_table_rows(lines):
    """Return list of lists of cell strings (skip separator row)."""
    rows = []
    for line in lines:
        if is_separator_row(line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def strip_emoji(text):
    """Remove common emoji for Word compatibility."""
    emoji_map = {
        '✅': '[OK]', '❌': '[FAIL]', '⚡': '[Fast]', '👤': '',
        '🟢': 'Easy', '🟡': 'Med', '🔴': 'Hard', '🧮': '[Calc]',
        '📋': '', '📖': '', '📄': '', '🔢': '', '📝': '', '📊': '',
        '⏳': '[Pending]', '→': '→', '←': '←', '╎': '|',
    }
    for e, r in emoji_map.items():
        text = text.replace(e, r)
    # Remove remaining emoji (U+1F000+)
    text = re.sub(r'[\U0001F000-\U0001FFFF]', '', text)
    return text


def strip_md_links(text):
    """Remove markdown [text](url) → text, and bare #anchor refs."""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def add_md_table(doc, table_lines, is_appendix_b=False):
    """Render markdown table as a Word table, auto-sized to page width."""
    rows = parse_md_table_rows(table_lines)
    if not rows:
        return

    num_cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Available width (A4 - margins = 21 - 3 - 2 = 16 cm)
    total_cm = 16.0

    # Column width strategies per table type
    if is_appendix_b and num_cols == 4:
        widths_cm = [1.5, 1.2, 10.3, 3.0]
    elif num_cols == 2:
        widths_cm = [5.5, 10.5]
    elif num_cols == 3:
        widths_cm = [4.0, 8.5, 3.5]
    elif num_cols == 4:
        widths_cm = [3.0, 5.0, 5.5, 2.5]
    elif num_cols == 5:
        widths_cm = [2.0, 4.0, 4.0, 3.5, 2.5]
    elif num_cols == 6:
        widths_cm = [1.5, 3.5, 3.5, 3.0, 2.0, 2.5]
    else:
        per = total_cm / num_cols
        widths_cm = [per] * num_cols

    # Scale if over total
    total_set = sum(widths_cm)
    if total_set > total_cm:
        factor = total_cm / total_set
        widths_cm = [w * factor for w in widths_cm]

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)

        for c_idx, cell_text in enumerate(row_data[:num_cols]):
            cell = row.cells[c_idx]
            try:
                cell.width = Cm(widths_cm[c_idx])
            except Exception:
                pass
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if is_header:
                shade_cell(cell, "2E74B5")

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

            text = strip_emoji(cell_text)
            run = p.add_run(text)
            if is_header:
                set_font_vn(run, bold=True, size=SIZE_TABLE,
                            color=RGBColor(0xFF, 0xFF, 0xFF))
            else:
                set_font_vn(run, size=SIZE_TABLE)
                # Alternate row shading
                if r_idx % 2 == 0:
                    shade_cell(cell, "EBF3FB")

    # Space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── Code block ────────────────────────────────────────────────────
def add_code_block(doc, lines_list):
    """Render code/ASCII as a styled box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)

    # Gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    # Join lines and add as single run
    code_text = '\n'.join(lines_list)
    run = p.add_run(code_text)
    run.font.name = FONT_CODE
    run.font.size = SIZE_CODE
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Bullet list ────────────────────────────────────────────────────
def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.8 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after   = Pt(3)

    bullet_char = "•" if level == 0 else "–"
    run_bullet = p.add_run(bullet_char + "  ")
    set_font_vn(run_bullet, bold=(level == 0), size=SIZE_BODY)
    # text already has leading '- ' stripped by regex; just strip whitespace
    add_inline(p, strip_emoji(text.strip()))
    return p


# ─── Main parser ────────────────────────────────────────────────────
def parse_and_build(doc: Document):
    with open(SRC, encoding='utf-8') as f:
        raw_lines = f.readlines()

    lines = [l.rstrip('\n') for l in raw_lines]
    i = 0
    in_code = False
    code_buf = []
    table_buf = []
    in_table = False
    is_appendix_b_table = False
    in_toc = False  # skip TOC section

    def flush_table():
        nonlocal table_buf, in_table, is_appendix_b_table
        if table_buf:
            add_md_table(doc, table_buf, is_appendix_b=is_appendix_b_table)
        table_buf = []
        in_table = False
        is_appendix_b_table = False

    while i < len(lines):
        line = lines[i]

        # ── Code block toggle ──────────────────────────────────────
        if line.strip().startswith('```'):
            if in_code:
                # End code block
                if code_buf and not code_buf[0].strip().lower().startswith('mermaid'):
                    add_code_block(doc, code_buf)
                elif code_buf:
                    # Mermaid block — render as note
                    p = doc.add_paragraph()
                    r = p.add_run("[ Xem sơ đồ tại file BAO_CAO_DU_AN_v2.md — Mermaid flowchart ]")
                    set_font_vn(r, italic=True, size=Pt(10), color=RGBColor(0x80, 0x80, 0x80))
                    p.paragraph_format.left_indent = Cm(1)
                in_code = False
                code_buf = []
            else:
                if in_table:
                    flush_table()
                in_code = True
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Table detection ────────────────────────────────────────
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                # Detect appendix B tables
                # Look back for context
                for back in range(max(0, i-5), i):
                    if lines[back].startswith('### B.'):
                        is_appendix_b_table = True
                        break
            table_buf.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        stripped = line.strip()

        # ── Skip horizontal rules & empty ─────────────────────────
        if stripped == '' or stripped == '---' or stripped.startswith('> **Lưu ý:**') is False and re.match(r'^-{3,}$', stripped):
            if stripped == '':
                pass
            elif re.match(r'^-{3,}$', stripped):
                # HR — add some spacing
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Headings ───────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading_text = strip_emoji(strip_md_links(m.group(2).strip()))
            # Detect and skip TOC section (## MỤC LỤC)
            if 'MỤC LỤC' in heading_text.upper():
                in_toc = True
                i += 1
                continue
            elif in_toc and level <= 2:
                # First real section after TOC ends TOC mode
                in_toc = False
            if level == 1:
                add_h1(doc, heading_text)
            elif level == 2:
                add_h2(doc, heading_text)
            elif level == 3:
                add_h3(doc, heading_text)
            else:
                add_h4(doc, heading_text)
            i += 1
            continue

        # Skip TOC content
        if in_toc:
            i += 1
            continue

        # ── Bullet list ────────────────────────────────────────────
        m_bullet = re.match(r'^(\s*)-\s+(.*)', line)
        if m_bullet:
            indent_spaces = len(m_bullet.group(1))
            level = indent_spaces // 2
            bullet_text = strip_md_links(m_bullet.group(2))
            # Skip pure TOC nav links: [text](#anchor)
            if re.match(r'^\[.+\]\(#', m_bullet.group(2).strip()):
                i += 1
                continue
            add_bullet(doc, bullet_text, level=level)
            i += 1
            continue

        # Numbered list: 1. 2. 3.
        m_num = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if m_num:
            indent_spaces = len(m_num.group(1))
            level = indent_spaces // 2
            num_text = strip_md_links(m_num.group(2))
            # Skip TOC links
            if re.match(r'^\[.+\]\(#', m_num.group(2).strip()):
                i += 1
                continue
            add_bullet(doc, num_text, level=level)
            i += 1
            continue

        # ── Blockquote / note ──────────────────────────────────────
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(strip_emoji(strip_md_links(stripped.lstrip('> '))))
            set_font_vn(run, italic=True, size=Pt(11), color=RGBColor(0x60, 0x60, 0x60))
            i += 1
            continue

        # ── Skip remaining TOC-style nav links ─────────────────────
        if re.match(r'^\d+\.\s+\[', stripped) and '#' in stripped:
            i += 1
            continue

        # ── Normal paragraph ───────────────────────────────────────
        if stripped:
            clean = strip_emoji(strip_md_links(stripped))
            # Remove image links
            clean = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[\1]', clean)
            add_body(doc, clean)

        i += 1

    # Flush any remaining table
    if in_table and table_buf:
        flush_table()


# ─── Main ──────────────────────────────────────────────────────────
def main():
    print("⏳ Building professional DOCX report...")
    doc = create_doc()

    # Header
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TaxAI — Hệ thống Tư vấn Thuế Tự động  |  v2.0  |  10/04/2026")
    run.font.name = FONT_MAIN
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True

    # Footer with page numbers
    add_page_number(doc)

    # Cover page
    add_cover(doc)

    # Parse and add content
    parse_and_build(doc)

    doc.save(DST)
    print(f"✅ Saved: {DST}")
    print(f"   File size: {DST.stat().st_size / 1024:.0f} KB")


if __name__ == '__main__':
    main()
