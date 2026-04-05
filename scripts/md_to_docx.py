"""
Convert BAO_CAO_DU_AN.md to .docx using python-docx.
Usage: python scripts/md_to_docx.py
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_line(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def apply_inline_formatting(paragraph, text):
    """Apply bold, italic, inline code to a paragraph run."""
    # Split by bold/italic/code markers
    # Handle **bold**, *italic*, `code`
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            if part:
                paragraph.add_run(part)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (table_data, next_idx)."""
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        elif re.match(r'^\s*\|?[-:| ]+\|?\s*$', line) and '---' in line:
            i += 1  # skip separator
        else:
            break

    if not table_lines:
        return None, start_idx + 1

    # Parse header and data rows
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    return rows, i


def md_to_docx(md_path: str, docx_path: str):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Set heading styles
    for i in range(1, 5):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Times New Roman'
        h.font.bold = True
        sizes = {1: 16, 2: 14, 3: 13, 4: 12}
        h.font.size = Pt(sizes[i])
        h.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # End code block — add as paragraph with code style
                in_code_block = False
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                # Add light gray shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F0F0')
                pPr.append(shd)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___') or re.match(r'^-{3,}$', stripped):
            add_horizontal_line(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # Remove markdown links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            text = re.sub(r'[*_`]', '', text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and stripped.endswith('|'):
            # Collect all table lines
            table_lines = []
            header_row = None
            j = i
            while j < len(lines):
                tline = lines[j].strip()
                if not tline:
                    break
                if tline.startswith('|') and tline.endswith('|'):
                    cells = [c.strip() for c in tline.strip('|').split('|')]
                    table_lines.append(cells)
                elif re.match(r'^\|[-: ]+\|', tline) or re.match(r'^[-: |]+$', tline):
                    # separator row — skip
                    pass
                else:
                    break
                j += 1

            if table_lines:
                num_cols = max(len(r) for r in table_lines)
                # Filter out separator rows
                data_rows = [r for r in table_lines if not all(
                    re.match(r'^[-: ]+$', c) for c in r
                )]

                if data_rows:
                    tbl = doc.add_table(rows=len(data_rows), cols=num_cols)
                    tbl.style = 'Table Grid'

                    for ri, row_data in enumerate(data_rows):
                        row = tbl.rows[ri]
                        for ci, cell_text in enumerate(row_data[:num_cols]):
                            if ci < len(row.cells):
                                cell = row.cells[ci]
                                p = cell.paragraphs[0]
                                p.clear()
                                # Header row styling
                                if ri == 0:
                                    set_cell_bg(cell, '1A5276')
                                    run = p.add_run(re.sub(r'[*_`]', '', cell_text))
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    run.font.size = Pt(10)
                                else:
                                    cell_text_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                                    run = p.add_run(re.sub(r'[*_]', '', cell_text_clean))
                                    run.font.size = Pt(10)
                                    if ri % 2 == 0:
                                        set_cell_bg(cell, 'EBF5FB')
                doc.add_paragraph()  # spacing after table
            i = j
            continue

        # Blockquote (> text)
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '8')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '2874A6')
            pBdr.append(left)
            pPr.append(pBdr)
            apply_inline_formatting(p, text)
            p.runs[0].font.color.rgb = RGBColor(0x27, 0x6D, 0xA0) if p.runs else None
            i += 1
            continue

        # Bullet list
        if re.match(r'^[-*+]\s+', stripped):
            text = re.sub(r'^[-*+]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            apply_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            apply_inline_formatting(p, text)
            i += 1
            continue

        # Bold/strong Q-number lines (Q1, Q2, etc.)
        if re.match(r'^\*\*Q\d+\*\*', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            apply_inline_formatting(p, stripped)
            i += 1
            continue

        # Normal paragraph (with inline formatting)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        apply_inline_formatting(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")


if __name__ == '__main__':
    md_path = 'BAO_CAO_DU_AN.md'
    docx_path = 'BAO_CAO_DU_AN.docx'
    md_to_docx(md_path, docx_path)
    print(f"Done! File size: {Path(docx_path).stat().st_size / 1024:.0f} KB")
