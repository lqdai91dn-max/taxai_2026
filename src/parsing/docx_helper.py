"""
DocxHelper — Extract text và tables từ file Word (.docx hoặc .doc).

Trả về cùng signature với SmartPDFHelper.extract_text_and_tables() để
pipeline.py có thể dùng đồng nhất mà không cần thay đổi Stage 2-5.

Output text giữ nguyên từng paragraph trên 1 dòng → state machine
nhận dạng Chương/Điều/Khoản chính xác như với PDF digital.

.doc (OLE2): dùng antiword với UTF-8 mapping để extract text + parse bảng
từ định dạng |cell1|cell2|.
"""

import subprocess
import re
from pathlib import Path
from typing import Tuple, List, Dict, Any

import docx

try:
    from ..utils.logger import logger
except Exception:
    import logging
    logger = logging.getLogger(__name__)


def extract_text_and_tables_from_docx(
    docx_path: Path,
) -> Tuple[str, List[Dict[str, Any]], int, Dict[str, Any]]:
    """
    Extract text và tables từ .docx file.

    Returns:
        (text, tables, total_pages, metadata)
        — cùng signature với SmartPDFHelper.extract_text_and_tables()

    Notes:
        - total_pages = 0 (Word không expose page info qua python-docx API)
        - table page_number = 0 vì cùng lý do trên
        - Text từ table cells KHÔNG lẫn vào full_text (truy cập qua doc.tables riêng)
    """
    logger.info(f"📝 DocxHelper: extracting {docx_path.name}")

    doc = docx.Document(str(docx_path))

    # ── Text extraction ───────────────────────────────────────────────────
    # doc.paragraphs chỉ trả về paragraph trong body, không bao gồm cell trong bảng
    # Regex: "Điều N " without period (e.g. "Điều 29 Title") → normalize to "Điều 29. Title"
    _RE_DIEU_NO_PERIOD = re.compile(r'^(Điều\s+\d+[a-zA-Z]*)\s+([^\s])', re.UNICODE)

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Normalize heading paragraphs that are missing the period after article number
        if 'heading' in para.style.name.lower():
            text = _RE_DIEU_NO_PERIOD.sub(r'\1. \2', text)
        lines.append(text)

    full_text = "\n".join(lines)

    # ── Table extraction ──────────────────────────────────────────────────
    tables_data = []
    for idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        raw_rows = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            raw_rows.append(cells)

        if not raw_rows:
            continue

        headers = raw_rows[0]
        data_rows = [r for r in raw_rows[1:] if any(c for c in r)]

        # Bỏ qua bảng rỗng (chỉ có header, không có data)
        if not data_rows:
            continue

        tables_data.append({
            "table_index": idx,
            "headers": headers,
            "rows": data_rows,
            "row_count": len(data_rows),
            "col_count": len(headers),
            "page_number": 0,
            "extraction_strategy": "docx",
        })

    metadata = {
        "pdf_type": "docx",
        "total_pages": 0,
        "text_length": len(full_text),
        "tables_found": len(tables_data),
        "extraction_method": "python-docx",
    }

    logger.info(
        f"✅ DocxHelper done: {len(full_text)} chars, "
        f"{len(tables_data)} tables"
    )

    return full_text, tables_data, 0, metadata


def extract_text_and_tables_from_doc(
    doc_path: Path,
) -> Tuple[str, List[Dict[str, Any]], int, Dict[str, Any]]:
    """
    Extract text và tables từ .doc (OLE2 binary) file qua antiword.

    antiword -m UTF-8.txt → text với Vietnamese diacritics đúng.
    Bảng từ antiword dùng định dạng |cell1|cell2|... → parse thành structured tables.

    Returns:
        (text, tables, total_pages, metadata)  — cùng signature với docx helper.
    """
    logger.info(f"📝 DocHelper (.doc): extracting {doc_path.name} via antiword")

    result = subprocess.run(
        ["antiword", "-m", "UTF-8.txt", str(doc_path)],
        capture_output=True,
        encoding="utf-8",
        errors="replace",
    )
    if result.returncode != 0:
        raise RuntimeError(f"antiword failed on {doc_path.name}: {result.stderr[:200]}")

    raw = result.stdout

    # ── Parse tables + text ───────────────────────────────────────────────
    # antiword renders tables as runs of lines matching |cell|cell|...|
    # Non-table lines are plain text paragraphs.
    text_lines: List[str] = []
    tables_data: List[Dict[str, Any]] = []

    lines = raw.splitlines()
    i = 0
    table_idx = 0
    while i < len(lines):
        line = lines[i]
        # Detect table block: line starts and ends with |, contains at least one |
        if line.startswith("|") and line.count("|") >= 2:
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                # Skip separator lines (only -, |, spaces)
                if re.match(r"^\|[-|\s]+\|$", lines[i]):
                    i += 1
                    continue
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                # Parse each row: split on | and strip
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split("|") if c.strip()]
                    if cells:
                        rows.append(cells)
                if rows:
                    # Normalize column count to max across rows
                    max_cols = max(len(r) for r in rows)
                    rows = [r + [""] * (max_cols - len(r)) for r in rows]
                    headers = rows[0]
                    data_rows = [r for r in rows[1:] if any(c for c in r)]
                    if data_rows:
                        tables_data.append({
                            "table_index": table_idx,
                            "headers": headers,
                            "rows": data_rows,
                            "row_count": len(data_rows),
                            "col_count": max_cols,
                            "page_number": 0,
                            "extraction_strategy": "antiword",
                        })
                        table_idx += 1
        else:
            stripped = line.strip()
            if stripped:
                text_lines.append(stripped)
            i += 1

    full_text = "\n".join(text_lines)

    metadata = {
        "pdf_type": "doc",
        "total_pages": 0,
        "text_length": len(full_text),
        "tables_found": len(tables_data),
        "extraction_method": "antiword",
    }

    logger.info(
        f"✅ DocHelper done: {len(full_text)} chars, "
        f"{len(tables_data)} tables"
    )

    return full_text, tables_data, 0, metadata
